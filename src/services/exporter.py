"""
Módulo de Exportación DOCX y PDF para ActaClara (P4/P5 - v0.6).

Provee las clases ``DocxExporter`` y ``PdfExporter``, y la dataclass
``ActaMetadata`` para generar documentos corporativos a partir de
actas transcritas y normalizadas.

Dependencias:
    - python-docx
    - fpdf2  (``pip install fpdf2``)
    - src.models.acta.Acta
    - src.database.db_manager.DBManager
"""

from __future__ import annotations

import logging
import sqlite3
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH # type: ignore
from docx.shared import Pt, RGBColor # type: ignore
from docx.oxml import OxmlElement # type: ignore
from docx.oxml.ns import qn # type: ignore

from src.database.db_manager import DBManager # type: ignore
from src.models.acta import Acta # type: ignore

logger = logging.getLogger("DocxExporter")


# ---------------------------------------------------------------------------
# Dataclass: Metadatos del Acta
# ---------------------------------------------------------------------------

@dataclass
class ActaMetadata:
    """Metadatos complementarios capturados en la UI para enriquecer el DOCX.

    Attributes:
        proyecto: Nombre del proyecto asociado a la reunión.
        objetivo: Objetivo principal de la reunión.
        asistentes: Lista de nombres de los participantes.
        acuerdos: Lista de acuerdos alcanzados durante la reunión.
        compromisos: Lista de compromisos adquiridos con responsable y plazo.
    """

    proyecto: str = ""
    objetivo: str = ""
    asistentes: list[str] = field(default_factory=list)
    acuerdos: list[str] = field(default_factory=list)
    compromisos: list[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Interfaz abstracta (ISP – Interface Segregation Principle)
# ---------------------------------------------------------------------------

class ExporterInterface:
    """Contrato base para cualquier exportador de actas.

    Define los métodos que todo exportador concreto debe implementar,
    permitiendo extender a PDF, HTML, etc. sin modificar ``DocxExporter``
    (OCP – Open/Closed Principle).
    """

    def apply_template(self) -> None:
        """Configura la plantilla y estilos corporativos del documento."""
        raise NotImplementedError

    def add_section(self, title: str, content: str) -> None:
        """Agrega una sección con título y contenido al documento.

        Args:
            title: Encabezado de la sección (e.g. "Asistentes").
            content: Cuerpo de texto de la sección.
        """
        raise NotImplementedError

    def save_and_record(self, acta_id: int, output_path: Path) -> Path:
        """Persiste el documento y registra la ruta en la base de datos.

        Args:
            acta_id: Identificador único del acta en SQLite.
            output_path: Ruta destino del archivo generado.

        Returns:
            La ruta final del archivo guardado.
        """
        raise NotImplementedError


# ---------------------------------------------------------------------------
# Implementación concreta: Exportador DOCX
# ---------------------------------------------------------------------------

class DocxExporter(ExporterInterface):
    """Exportador de actas al formato Microsoft Word (.docx).

    Responsabilidades (SRP – Single Responsibility Principle):
        - Construir un ``Document`` de *python-docx* a partir de un ``Acta``
          y sus ``ActaMetadata``.
        - Aplicar estilos corporativos definidos en ``apply_template``.
        - Persistir el archivo y notificar al ``DBManager``.

    Attributes:
        acta: Objeto ``Acta`` con la transcripción normalizada.
        metadata: Metadatos complementarios capturados en la interfaz.
        document: Instancia de ``Document`` de *python-docx*.
        db_manager: Referencia al gestor de base de datos (DIP – Dependency
            Inversion Principle: se inyecta, no se instancia internamente).

    Example:
        >>> exporter = DocxExporter(acta=mi_acta, metadata=meta)
        >>> exporter.apply_template()
        >>> exporter.add_section("Asistentes", "Juan, María, Pedro")
        >>> path = exporter.save_and_record(acta_id=1, output_path=Path("out.docx"))
    """

    _COLOR_TITULO: RGBColor
    _COLOR_SECCION: RGBColor
    _FUENTE_BASE: str
    _ALIGN_TEXT: int

    def __init__(
        self,
        acta: Acta,
        metadata: ActaMetadata,
        db_manager: Optional[DBManager] = None,
        template_name: str = "Corporativa Formal",
    ) -> None:
        """Inicializa el exportador con los datos del acta.

        Args:
            acta: Objeto ``Acta`` que contiene el texto ya normalizado.
            metadata: Información complementaria capturada en la UI.
            db_manager: Gestor de BD. Si es ``None`` se usa la instancia
                Singleton de ``DBManager``.
            template_name: Nombre del tipo de plantilla a aplicar ("Corporativa Formal", "Académica", "Minimalista").
        """
        self.acta: Acta = acta
        self.metadata: ActaMetadata = metadata
        self.document: Document = Document()
        self.db_manager: DBManager = db_manager or DBManager()
        self.template_name = template_name
        self._apply_template_config()

    def _apply_template_config(self) -> None:
        """Configura dinámicamente los estilos base según la plantilla escogida."""
        if self.template_name == "Académica":
            self._COLOR_TITULO: RGBColor = RGBColor(0x00, 0x00, 0x00)
            self._COLOR_SECCION: RGBColor = RGBColor(0x00, 0x00, 0x00)
            self._FUENTE_BASE: str = "Times New Roman"
            self._ALIGN_TEXT = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif self.template_name == "Minimalista":
            self._COLOR_TITULO = RGBColor(0x21, 0x21, 0x21)
            self._COLOR_SECCION = RGBColor(0x64, 0x64, 0x64)
            self._FUENTE_BASE = "Arial"
            self._ALIGN_TEXT = WD_ALIGN_PARAGRAPH.LEFT
        else: # Corporativa Formal u otro fallback
            self._COLOR_TITULO = RGBColor(0x1F, 0x49, 0x7D)
            self._COLOR_SECCION = RGBColor(0x2E, 0x74, 0xB5)
            self._FUENTE_BASE = "Calibri"
            self._ALIGN_TEXT = WD_ALIGN_PARAGRAPH.JUSTIFY

    # -- Métodos públicos ---------------------------------------------------

    def apply_template(self) -> None:
        """Configura la plantilla corporativa del documento DOCX.

        Establece fuentes base del documento y agrega el título principal
        ``ACTA DE REUNIÓN`` con estilo corporativo centrado. También
        construye encabezado y pie de página.
        """
        # Configurar estilos globales del documento
        style_normal = self.document.styles["Normal"]
        style_normal.font.name = self._FUENTE_BASE
        style_normal.font.size = Pt(11)

        # Configurar estilo de título
        style_titulo = self.document.styles["Title"]
        style_titulo.font.name = self._FUENTE_BASE
        style_titulo.font.size = Pt(20)
        style_titulo.font.bold = True
        style_titulo.font.color.rgb = self._COLOR_TITULO

        # Título principal del documento
        titulo_par = self.document.add_heading("ACTA DE REUNIÓN", level=0)
        titulo_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Asegurar fuente y color correctos en el run generado
        for run in titulo_par.runs:
            run.font.name = self._FUENTE_BASE
            run.font.color.rgb = self._COLOR_TITULO

        # Subtítulo: nombre del proyecto
        if self.metadata.proyecto:
            sub = self.document.add_paragraph(self.metadata.proyecto)
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_run = sub.runs[0]
            sub_run.font.name = self._FUENTE_BASE
            sub_run.font.size = Pt(13)
            sub_run.font.italic = True
            sub_run.font.color.rgb = self._COLOR_SECCION

        # Línea separadora horizontal
        self.document.add_paragraph("─" * 60)

        self._build_header()
        self._build_footer()

        logger.info(
            "Plantilla corporativa aplicada al documento para acta '%s'.",
            self.acta.titulo,
        )

    def add_section(self, title: str, content: str) -> None:
        """Agrega una sección formateada al documento.

        Añade un encabezado de nivel 2 con el ``title`` en negrita y
        color corporativo, seguido de un párrafo con el ``content``.

        Args:
            title: Título de la sección visible en el documento.
            content: Cuerpo de texto de la sección.
        """
        # Encabezado de sección (Heading 2)
        heading = self.document.add_heading(title, level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in heading.runs:
            run.font.name = self._FUENTE_BASE
            run.font.bold = True
            run.font.color.rgb = self._COLOR_SECCION

        # Limpiar timestamps del contenido [00:00] antes de exportar
        clean_content = re.sub(r'\[\d{1,2}:\d{2}\]\s*', '', content)

        # Contenido de la sección
        parrafo = self.document.add_paragraph(clean_content)
        parrafo.alignment = self._ALIGN_TEXT
        for run in parrafo.runs:
            run.font.name = self._FUENTE_BASE
            run.font.size = Pt(11)

        # Espacio posterior a cada sección
        self.document.add_paragraph()

    def save_and_record(self, acta_id: int, output_path: Path) -> Path:
        """Guarda el DOCX en disco y actualiza la BD con la ruta generada.

        Crea los directorios intermedios si no existen. Persiste el archivo
        y ejecuta un ``UPDATE`` en la tabla ``actas`` para registrar la ruta.

        Args:
            acta_id: ID del acta en la tabla ``actas``.
            output_path: Ruta de destino del archivo ``.docx``.

        Returns:
            ``Path`` absoluto del archivo guardado exitosamente.

        Raises:
            IOError: Si no es posible escribir en ``output_path``.
            sqlite3.Error: Si falla la actualización en la BD.
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Persistir el documento en disco
        try:
            self.document.save(str(output_path))
            logger.info("Documento DOCX guardado en: %s", output_path)
        except OSError as exc:
            logger.error("No se pudo guardar el archivo DOCX: %s", exc)
            raise IOError(f"Error al guardar el DOCX en '{output_path}': {exc}") from exc

        # Registrar la ruta en la base de datos
        self._update_docx_ruta_in_db(acta_id, str(output_path))

        return output_path.resolve()

    # -- Métodos internos ---------------------------------------------------

    def _build_header(self) -> None:
        """Construye el encabezado de sección con metadatos informativos.

        Agrega una tabla de dos columnas con: Fecha de creación del acta
        y el objetivo de la reunión.
        """
        fecha_str = (
            self.acta.fecha_creacion.strftime("%d/%m/%Y %H:%M")
            if self.acta.fecha_creacion
            else datetime.now().strftime("%d/%m/%Y %H:%M")
        )

        # Tabla de metadatos de cabecera
        tabla = self.document.add_table(rows=2, cols=2)
        tabla.style = "Table Grid"

        # Fila 0: etiquetas
        tabla.cell(0, 0).text = "Fecha:"
        tabla.cell(0, 1).text = fecha_str
        # Fila 1: objetivo
        tabla.cell(1, 0).text = "Objetivo:"
        tabla.cell(1, 1).text = self.metadata.objetivo or "—"

        # Aplicar negrita a los labels
        for fila in range(2):
            celda_label = tabla.cell(fila, 0)
            for par in celda_label.paragraphs:
                for run in par.runs:
                    run.font.bold = True
        self.document.add_paragraph()

    def _build_footer(self) -> None:
        """Agrega información de versión y número de página en el pie de página del documento."""
        seccion = self.document.sections[0]
        footer = seccion.footer
        
        # Eliminar párrafos existentes si los hay para evitar duplicados
        for p in footer.paragraphs:
            p.text = ""

        footer_par = footer.paragraphs[0]
        footer_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 1. Info de versión (Lado izquierdo/centro)
        run_info = footer_par.add_run(
            f"ActaClara | Idioma: {self.acta.idioma} | "
            f"Diccionario v{self.acta.version_diccionario}"
        )
        run_info.font.size = Pt(8)
        run_info.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        # 2. Número de página (Lado derecho - usando tabuladores o alineación)
        # En Word, una forma limpia es añadir un campo PAGE
        footer_par.add_run("\t\tPágina ")
        self._add_page_number(footer_par)

    def _add_page_number(self, paragraph) -> None:
        """Inserta un campo XML de número de página en el párrafo."""
        run = paragraph.add_run()
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar)
        run._r.append(instrText)
        run._r.append(fldChar2)
    def _build_all_sections(self, texto_normalizado: str) -> None:
        """Orquesta la creación de todas las secciones del acta.

        Utiliza ``self.metadata`` y ``self.acta`` para poblar:

            1. Información General (Proyecto, Objetivo, Fecha).
            2. Asistentes.
            3. Desarrollo de la Reunión (texto normalizado).
            4. Acuerdos y Compromisos.

        Args:
            texto_normalizado: Transcripción ya procesada con los modismos
                sustituidos por sus equivalentes formales.
        """
        # 1. Información General
        fecha_str = (
            self.acta.fecha_creacion.strftime("%d/%m/%Y %H:%M")
            if self.acta.fecha_creacion
            else datetime.now().strftime("%d/%m/%Y %H:%M")
        )
        info_general = (
            f"Proyecto: {self.metadata.proyecto or '—'}\n"
            f"Objetivo: {self.metadata.objetivo or '—'}\n"
            f"Fecha: {fecha_str}"
        )
        self.add_section("1. Información General", info_general)

        # 2. Asistentes
        asistentes_texto = (
            "\n".join(f"• {nombre}" for nombre in self.metadata.asistentes)
            if self.metadata.asistentes
            else "Sin asistentes registrados."
        )
        self.add_section("2. Asistentes", asistentes_texto)

        # 3. Desarrollo de la Reunión
        self.add_section(
            "3. Desarrollo de la Reunión",
            texto_normalizado or "Sin transcripción disponible.",
        )

        # 4. Acuerdos
        acuerdos_texto = (
            "\n".join(f"• {acuerdo}" for acuerdo in self.metadata.acuerdos)
            if self.metadata.acuerdos
            else "Sin acuerdos registrados."
        )
        self.add_section("4. Acuerdos", acuerdos_texto)

        # 5. Compromisos
        compromisos_texto = (
            "\n".join(f"• {compromiso}" for compromiso in self.metadata.compromisos)
            if self.metadata.compromisos
            else "Sin compromisos registrados."
        )
        self.add_section("5. Compromisos", compromisos_texto)

    def _update_docx_ruta_in_db(self, acta_id: int, ruta: str) -> None:
        """Ejecuta UPDATE en SQLite para registrar la ruta del DOCX generado.

        Args:
            acta_id: ID del registro en la tabla ``actas``.
            ruta: Ruta absoluta del archivo ``.docx`` generado.

        Raises:
            sqlite3.Error: Si falla la operación SQL.
        """
        query = "UPDATE actas SET archivo_docx_ruta = ? WHERE id = ?"
        try:
            with self.db_manager._get_connection() as conn:
                conn.execute(query, (ruta, acta_id))
            logger.info(
                "Ruta DOCX '%s' registrada en BD para acta ID=%d.", ruta, acta_id
            )
        except sqlite3.Error as exc:
            logger.error(
                "Error al actualizar archivo_docx_ruta en BD (acta_id=%d): %s",
                acta_id,
                exc,
            )
            raise


# ---------------------------------------------------------------------------
# Implementación concreta: Exportador PDF (fpdf2)
# ---------------------------------------------------------------------------

class PdfExporter(ExporterInterface):
    """Exportador de actas al formato PDF usando fpdf2.

    Responsabilidades:
        - Construir un documento PDF con encabezado, pie de página,
          tabla de metadatos y secciones de contenido.
        - Mantener la misma API pública que ``DocxExporter`` para que
          el código cliente pueda intercambiar formatos sin modificarse.

    Attributes:
        acta: Objeto ``Acta`` con la transcripción normalizada.
        metadata: Metadatos complementarios capturados en la interfaz.
        pdf: Instancia de ``fpdf2.FPDF``.
        db_manager: Referencia al gestor de base de datos (DIP).

    Example:
        >>> exporter = PdfExporter(acta=mi_acta, metadata=meta)
        >>> exporter.apply_template()
        >>> exporter.add_section("Asistentes", "Juan, María, Pedro")
        >>> path = exporter.save_and_record(acta_id=1, output_path=Path("out.pdf"))
    """

    _COLOR_TITULO: tuple
    _COLOR_SECCION: tuple
    _COLOR_GRIS: tuple
    _FUENTE_BASE: str
    _ALIGN_TEXT: str

    def __init__(
        self,
        acta: Acta,
        metadata: ActaMetadata,
        db_manager: Optional[DBManager] = None,
        template_name: str = "Corporativa Formal",
    ) -> None:
        """Inicializa el exportador PDF.

        Args:
            acta: Objeto ``Acta`` con el texto ya normalizado.
            metadata: Información complementaria capturada en la UI.
            db_manager: Gestor de BD. Si es ``None`` se usa la instancia
                Singleton de ``DBManager``.
            template_name: Nombre del tipo de plantilla a aplicar ("Corporativa Formal", "Académica", "Minimalista").
        """
        try:
            from fpdf import FPDF # type: ignore
        except ImportError as exc:
            raise ImportError(
                "Se requiere fpdf2 para PdfExporter. "
                "Instálalo con: pip install fpdf2"
            ) from exc

        self.acta: Acta = acta
        self.metadata: ActaMetadata = metadata
        self.db_manager: DBManager = db_manager or DBManager()
        
        # Clase interna para manejar el footer dinámico de fpdf2
        class CustomFPDF(FPDF):
            def __init__(self, exporter_ref, *args, **kwargs):
                super().__init__(*args, **kwargs)
                self.exporter = exporter_ref

            def footer(self):
                # Posición a 1.5 cm del final
                self.set_y(-15)
                self.set_font(self.exporter._FUENTE_BASE, "I", 8)
                self.set_text_color(128, 128, 128)
                
                # Texto de info (izquierda/centro)
                info = (f"ActaClara | Idioma: {self.exporter.acta.idioma} | "
                        f"Diccionario v{self.exporter.acta.version_diccionario}")
                
                # Número de página a la derecha
                page_num = f"Página {self.page_no()}"
                
                # Renderizar ambos en la misma línea
                self.cell(0, 10, self.exporter._clean_for_pdf(info), align="L")
                self.cell(0, 10, self.exporter._clean_for_pdf(page_num), align="R")

        self.pdf: FPDF = CustomFPDF(self, orientation="P", unit="mm", format="Letter")
        self.pdf.set_auto_page_break(auto=True, margin=25)
        
        self.template_name = template_name
        self._COLOR_GRIS = (128, 128, 128)
        self._apply_template_config()

    def _apply_template_config(self) -> None:
        """Configura dinámicamente los estilos base según la plantilla escogida en PDF."""
        if self.template_name == "Académica":
            self._COLOR_TITULO = (0, 0, 0)
            self._COLOR_SECCION = (0, 0, 0)
            self._FUENTE_BASE = "Times"
            self._ALIGN_TEXT = "J"
        elif self.template_name == "Minimalista":
            self._COLOR_TITULO = (33, 33, 33)
            self._COLOR_SECCION = (100, 100, 100)
            self._FUENTE_BASE = "Helvetica"
            self._ALIGN_TEXT = "L"
        else: # Corporativa Formal
            self._COLOR_TITULO = (31, 73, 125)
            self._COLOR_SECCION = (46, 116, 181)
            self._FUENTE_BASE = "Helvetica"
            self._ALIGN_TEXT = "J"

    def _clean_for_pdf(self, text: str) -> str:
        """Limpia caracteres Unicode que fpdf2 no soporta nativamente en sus fuentes core."""
        if not text: return ""
        replacements = {
            '\u2014': '--', # em-dash
            '\u2013': '-',  # en-dash
            '\u201c': '"',  # left double quote
            '\u201d': '"',  # right double quote
            '\u2018': "'",  # left single quote
            '\u2019': "'",  # right single quote
            '\u2026': '...',# ellipsis
            '\u2022': '-',  # bullets
        }
        for k, v in replacements.items():
            text = str(text).replace(k, v)
        # Reemplazar con '?' lo que quede fuera de latin-1
        return text.encode('latin-1', 'replace').decode('latin-1')

    # -- Métodos públicos ---------------------------------------------------

    def apply_template(self) -> None:
        """Configura la plantilla corporativa del documento PDF.

        Agrega la primera página, título ``ACTA DE REUNIÓN``, subtítulo
        con el nombre del proyecto, línea separadora y tabla de metadatos.
        También configura el pie de página automático.
        """
        self.pdf.add_page()

        # --- Título principal ---
        self.pdf.set_font(self._FUENTE_BASE, "B", 20)
        self.pdf.set_text_color(*self._COLOR_TITULO)
        self.pdf.cell(0, 12, self._clean_for_pdf("ACTA DE REUNIÓN"), new_x="LMARGIN", new_y="NEXT", align="C")

        # --- Subtítulo: proyecto ---
        if self.metadata.proyecto:
            self.pdf.set_font(self._FUENTE_BASE, "I", 13)
            self.pdf.set_text_color(*self._COLOR_SECCION)
            self.pdf.cell(
                0, 8, self._clean_for_pdf(self.metadata.proyecto),
                new_x="LMARGIN", new_y="NEXT", align="C",
            )

        # --- Línea separadora ---
        self.pdf.ln(3)
        y = self.pdf.get_y()
        self.pdf.set_draw_color(*self._COLOR_SECCION)
        self.pdf.set_line_width(0.5)
        self.pdf.line(10, y, self.pdf.w - 10, y)
        self.pdf.ln(5)

        self._build_header()

        logger.info(
            "Plantilla PDF corporativa aplicada para acta '%s'.",
            self.acta.titulo,
        )

    def add_section(self, title: str, content: str) -> None:
        """Agrega una sección formateada al PDF.

        Args:
            title: Título de la sección.
            content: Cuerpo de texto.
        """
        # Encabezado de sección
        self.pdf.set_font(self._FUENTE_BASE, "B", 13)
        self.pdf.set_text_color(*self._COLOR_SECCION)
        self.pdf.cell(0, 8, self._clean_for_pdf(title), new_x="LMARGIN", new_y="NEXT")

        # Limpiar timestamps del contenido [00:00] antes de exportar
        clean_content = re.sub(r'\[\d{1,2}:\d{2}\]\s*', '', content)

        # Contenido
        self.pdf.set_font(self._FUENTE_BASE, "", 11)
        self.pdf.set_text_color(0, 0, 0)
        self.pdf.multi_cell(0, 6, self._clean_for_pdf(clean_content), align=self._ALIGN_TEXT)
        self.pdf.ln(4)

    def save_and_record(self, acta_id: int, output_path: Path) -> Path:
        """Guarda el PDF en disco y actualiza la BD con la ruta generada.

        Args:
            acta_id: ID del acta en la tabla ``actas``.
            output_path: Ruta de destino del archivo ``.pdf``.

        Returns:
            ``Path`` absoluto del archivo guardado exitosamente.

        Raises:
            IOError: Si no es posible escribir en ``output_path``.
            sqlite3.Error: Si falla la actualización en la BD.
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # El pie de página se genera automáticamente en cada página gracias al override de CustomFPDF

        try:
            self.pdf.output(str(output_path))
            logger.info("Documento PDF guardado en: %s", output_path)
        except OSError as exc:
            logger.error("No se pudo guardar el archivo PDF: %s", exc)
            raise IOError(
                f"Error al guardar el PDF en '{output_path}': {exc}"
            ) from exc

        # Registrar en BD
        self._update_pdf_ruta_in_db(acta_id, str(output_path))

        return output_path.resolve()

    # -- Métodos internos ---------------------------------------------------

    def _build_header(self) -> None:
        """Tabla de metadatos: Fecha y Objetivo."""
        fecha_str = (
            self.acta.fecha_creacion.strftime("%d/%m/%Y %H:%M")
            if self.acta.fecha_creacion
            else datetime.now().strftime("%d/%m/%Y %H:%M")
        )

        col_w_label = 35
        col_w_value = self.pdf.w - 20 - col_w_label  # margen 10 + 10

        self.pdf.set_font(self._FUENTE_BASE, "B", 11)
        self.pdf.set_text_color(0, 0, 0)

        # Fila Fecha
        self.pdf.cell(col_w_label, 7, self._clean_for_pdf("Fecha:"), border=1)
        self.pdf.set_font(self._FUENTE_BASE, "", 11)
        self.pdf.cell(col_w_value, 7, self._clean_for_pdf(fecha_str), border=1, new_x="LMARGIN", new_y="NEXT")

        # Fila Objetivo
        self.pdf.set_font(self._FUENTE_BASE, "B", 11)
        self.pdf.cell(col_w_label, 7, self._clean_for_pdf("Objetivo:"), border=1)
        self.pdf.set_font(self._FUENTE_BASE, "", 11)
        objetivo = self.metadata.objetivo or "—"
        self.pdf.cell(col_w_value, 7, self._clean_for_pdf(objetivo), border=1, new_x="LMARGIN", new_y="NEXT")

        self.pdf.ln(6)

    def _build_footer(self) -> None:
        """El footer del PDF se maneja ahora vía el método footer() de la clase CustomFPDF."""
        pass

    def _build_all_sections(self, texto_normalizado: str) -> None:
        """Orquesta la creación de todas las secciones del acta en PDF.

        Args:
            texto_normalizado: Transcripción ya procesada.
        """
        # 1. Información General
        fecha_str = (
            self.acta.fecha_creacion.strftime("%d/%m/%Y %H:%M")
            if self.acta.fecha_creacion
            else datetime.now().strftime("%d/%m/%Y %H:%M")
        )
        info = (
            f"Proyecto: {self.metadata.proyecto or '—'}\n"
            f"Objetivo: {self.metadata.objetivo or '—'}\n"
            f"Fecha: {fecha_str}"
        )
        self.add_section("1. Información General", info)

        # 2. Asistentes
        asistentes = (
            "\n".join(f"• {n}" for n in self.metadata.asistentes)
            if self.metadata.asistentes
            else "Sin asistentes registrados."
        )
        self.add_section("2. Asistentes", asistentes)

        # 3. Desarrollo
        self.add_section(
            "3. Desarrollo de la Reunión",
            texto_normalizado or "Sin transcripción disponible.",
        )

        # 4. Acuerdos
        acuerdos = (
            "\n".join(f"• {a}" for a in self.metadata.acuerdos)
            if self.metadata.acuerdos
            else "Sin acuerdos registrados."
        )
        self.add_section("4. Acuerdos", acuerdos)

        # 5. Compromisos
        compromisos = (
            "\n".join(f"• {c}" for c in self.metadata.compromisos)
            if self.metadata.compromisos
            else "Sin compromisos registrados."
        )
        self.add_section("5. Compromisos", compromisos)

    def _update_pdf_ruta_in_db(self, acta_id: int, ruta: str) -> None:
        """Ejecuta UPDATE en SQLite para registrar la ruta del PDF generado.

        Reutiliza la columna ``archivo_docx_ruta`` (la tabla soporta una
        ruta de exportación genérica; si se requieren rutas separadas
        se ampliará el schema en un hito futuro).

        Args:
            acta_id: ID del registro en la tabla ``actas``.
            ruta: Ruta absoluta del archivo ``.pdf`` generado.
        """
        query = "UPDATE actas SET archivo_docx_ruta = ? WHERE id = ?"
        try:
            with self.db_manager._get_connection() as conn:
                conn.execute(query, (ruta, acta_id))
            logger.info(
                "Ruta PDF '%s' registrada en BD para acta ID=%d.", ruta, acta_id
            )
        except sqlite3.Error as exc:
            logger.error(
                "Error al actualizar archivo_docx_ruta (PDF) en BD (acta_id=%d): %s",
                acta_id,
                exc,
            )
            raise


# ---------------------------------------------------------------------------
# Factory: Selección de exportador por formato
# ---------------------------------------------------------------------------

def create_exporter(
    fmt: str,
    acta: Acta,
    metadata: ActaMetadata,
    db_manager: Optional[DBManager] = None,
    template_name: str = "Corporativa Formal",
) -> ExporterInterface:
    """Crea el exportador adecuado según el formato solicitado.

    Args:
        fmt: ``"docx"`` o ``"pdf"`` (case-insensitive).
        acta: Objeto ``Acta``.
        metadata: Metadatos complementarios.
        db_manager: Gestor de BD opcional.
        template_name: Nombre de la plantilla elegida.

    Returns:
        Instancia de ``DocxExporter`` o ``PdfExporter``.

    Raises:
        ValueError: Si *fmt* no es un formato soportado.
    """
    fmt_lower = fmt.strip().lower()
    if fmt_lower == "docx":
        return DocxExporter(acta=acta, metadata=metadata, db_manager=db_manager, template_name=template_name)
    if fmt_lower == "pdf":
        return PdfExporter(acta=acta, metadata=metadata, db_manager=db_manager, template_name=template_name)
    raise ValueError(
        f"Formato '{fmt}' no soportado. Use 'docx' o 'pdf'."
    )
